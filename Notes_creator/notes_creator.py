# to make error disappear follow the instructions to the file path and change
# "lis = BeautifulSoup(html).find_all('li')"    to    "lis = BeautifulSoup(html, features="html.parser").find_all('li')"
import wikipedia
import os
from docx import Document
from docx.shared import Pt
from datetime import date


def input_number(message, search_range):
    while True:
        try:
            user_input = int(input(message + '\n'))
        except ValueError:
            print("Not a number")
            continue
        else:
            if 0 < user_input <= search_range:
                return user_input
            else:
                print(f'out of range(0-{search_range})')


def add_definition(definition='', list_of_def=[]):
    if list_of_def is None:
        list_of_def = []
    if definition == '':
        return list_of_def
    list_of_def.append(definition)


print("Hello, with this program you can make automatically notes into your DOCX file")
lang_choice = input_number(f'choose language for your notes:'
                           f'\n1 - pl'
                           f'\n2 - eng', 2)
lang = "pl" if lang_choice == 1 else 'eng'
# sets lang for search on wikipedia
wikipedia.set_lang(lang)

while True:
    add_end = input_number('1 - add new note \n2 - save notes', 2)

    if add_end == 1:
        phrase = input('What you are looking for?(remember to use correct characters)')
        if phrase != '':
            phrase_results = wikipedia.search(phrase, results=5)

            if phrase_results:
                for i, phrase in enumerate(phrase_results, start=1):
                    print(f"{i} - {phrase}")

                option = input_number('which result interests you', 5) - 1

                length = input_number('Ok, You want:'
                                      '\n1 - only definition'
                                      '\n2 - full context', 2)
                try:
                    if length == 1:
                        definition = wikipedia.summary(phrase_results[option])
                        # saving only firsts lines of definition
                        text = definition[: definition.find('\n')]
                    elif length == 2:
                        text = wikipedia.summary(phrase_results[option])

                    # separate text to key of definition and definition
                    key_word = text[: text.find('–')]
                    definition = text[text.find('–'):]
                    # adds result to the list
                    add_definition([key_word, definition])
                    print(f'"{phrase_results[option]}" ADDED')
                except wikipedia.exceptions.DisambiguationError:
                    # if there is some wrong letter characters in search
                    print('You wrote something wrong! Remember to use: "ę,ć,ł..." etc.')
            else:
                # if search return nothing
                print(f'There is no such thing  as "{phrase}"')
        else:
            print('Write what you are looking for\n')
            continue

    elif add_end == 2:

        format_def = input_number("The notes you made are for(depends of your choice format of file will change):"
                                  "\n1 - learn"
                                  "\n2 - cheat", 2)
        document = Document()

        if format_def == 1:
            # saves notes with heading
            document.add_heading(f'Notes: {str(date.today())}\n', )
            p = document.add_paragraph()
            # resizing margins for docx file
            paragraph_format = p.paragraph_format
            paragraph_format.right_indent = Pt(-50)
            paragraph_format.left_indent = Pt(-50)

            for d in add_definition():
                # bolds key of definition
                p.add_run(d[0]).bold = True
                p.add_run(d[1] + '\n\n')

        elif format_def == 2:
            p = document.add_paragraph()
            # changes mright margin to fit "cheat" notes better
            paragraph_format = p.paragraph_format
            paragraph_format.right_indent = Pt(260)

            for d in add_definition():
                # resizing text
                key = p.add_run(d[0])
                key.bold = True
                key_size = key.font
                key_size.size = Pt(6)
                def_size = p.add_run(d[1] + '\n').font
                def_size.size = Pt(5)

        filename = f"Notes-{os.getlogin()}-{str(date.today())}"
        document.save(filename + '.docx')
        print(f'You can find your notes here:\n{os.getcwd()}\\{filename}.docx')
        input()
